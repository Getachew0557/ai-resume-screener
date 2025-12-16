import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract
import io
import logging
from typing import Union
from pathlib import Path
from backend.services.gemini_service import gemini_service

logger = logging.getLogger(__name__)

class DocumentParser:
    def extract_text(self, file_path: Union[str, Path], file_content: bytes = None) -> str:
        """
        Extracts text from PDF, DOCX, or Image files.
        If file_content is provided, it processes bytes directly (useful for API uploads).
        """
        file_path = str(file_path).lower()
        
        try:
            if file_path.endswith('.pdf'):
                return self._parse_pdf(file_path, file_content)
            elif file_path.endswith('.docx'):
                return self._parse_docx(file_path, file_content)
            elif file_path.endswith(('.png', '.jpg', '.jpeg', '.tiff')):
                return self._parse_image(file_path, file_content)
            elif file_path.endswith('.txt'):
                if file_content:
                    return file_content.decode('utf-8')
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError("Unsupported file format")
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise

    def _parse_pdf(self, file_path: str, file_content: bytes) -> str:
        doc = fitz.open(stream=file_content, filetype="pdf") if file_content else fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        
        # If text is very sparse, it might be an image scan. Use Gemini or OCR.
        if len(text.strip()) < 50: 
             # Fallback to Gemini for scanned PDFs if possible, or OCR. 
             # For now, let's try basic OCR or hint extracting images from PDF.
             # A robust solution would pass the PDF pages as images to Gemini.
             pass
        return text

    def _parse_docx(self, file_path: str, file_content: bytes) -> str:
        if file_content:
            doc = docx.Document(io.BytesIO(file_content))
        else:
            doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def _parse_image(self, file_path: str, file_content: bytes) -> str:
        try:
            # Prepare image for Gemini
            mime_type = "image/jpeg" # Defaulting to jpeg for simplicity, logical improvement would detect mime
            if file_path.lower().endswith(".png"):
                 mime_type = "image/png"
            
            if file_content:
                image_parts = [{"mime_type": mime_type, "data": file_content}]
            else:
                 with open(file_path, "rb") as f:
                     data = f.read()
                 image_parts = [{"mime_type": mime_type, "data": data}]

            prompt = "Extract all text from this image faithfully. Output only the text content."
            return gemini_service.generate_text(prompt, images=image_parts)

        except Exception as e:
            logger.warning(f"Gemini image parsing failed, falling back to Tesseract: {e}")
            # Fallback to Tesseract
            image = Image.open(io.BytesIO(file_content)) if file_content else Image.open(file_path)
            return pytesseract.image_to_string(image)

parser = DocumentParser()
